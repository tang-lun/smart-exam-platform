from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session

from app.db.database import get_db
from app.models.exam import Exam
from app.models.question import Question
from app.models.user import User
from app.models.exam_result import ExamResult
from app.schemas.exam import (
    ExamCreateRequest,
    ExamDetailResponse,
    ExamListResponse,
    ExamResponse,
    ExamResultListResponse,
    ExamResultResponse,
    ExamSubmitRequest,
)
from app.schemas.question import QuestionResponse
from app.services.ai_service import analyze_exam
from app.services.auth_service import get_current_user
from app.services.exam_service import allocate_question_scores, auto_select_questions

router = APIRouter(prefix="/api/exams", tags=["exams"])


def _exam_visible(user: User | None):
    if user is None:
        return Exam.owner_id == None  # noqa: E711
    return Exam.owner_id == user.id


def _infer_distribution(questions: list) -> dict[str, int]:
    """从已选题目反推题型分布。"""
    from collections import Counter
    types = Counter(q.type.value if hasattr(q.type, 'value') else str(q.type) for q in questions)
    return dict(types)


def _infer_diff_distribution(questions: list) -> dict[str, int]:
    """从已选题目反推难度分布（百分比）。"""
    from collections import Counter
    diffs = Counter(q.difficulty.value if hasattr(q.difficulty, 'value') else str(q.difficulty) for q in questions)
    total = sum(diffs.values())
    if total == 0:
        return {"easy": 30, "medium": 40, "hard": 30}
    return {k: round(v / total * 100) for k, v in diffs.items()}


@router.post("", response_model=ExamDetailResponse, status_code=201)
def create_exam(
    req: ExamCreateRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """创建试卷（支持手动选题或 AI 自动组卷）。"""
    if req.question_ids:
        questions = db.query(Question).filter(Question.id.in_(req.question_ids)).all()
        if len(questions) != len(req.question_ids):
            raise HTTPException(status_code=400, detail="部分题目不存在")
        question_ids = req.question_ids
    else:
        selected = auto_select_questions(
            db=db,
            knowledge_points=req.knowledge_points,
            difficulty_distribution=req.difficulty_distribution,
            type_distribution=req.type_distribution,
            grade_levels=req.grade_levels or None,
            owner_id=current_user.id,
        )
        if not selected:
            raise HTTPException(status_code=400, detail="题库中没有符合条件的题目，请先生成题目")
        question_ids = [q.id for q in selected]

    questions = db.query(Question).filter(Question.id.in_(question_ids)).all()
    q_map = {q.id: q for q in questions}
    ordered = [q_map[qid] for qid in question_ids if qid in q_map]

    # 自动分配每题分值（参考中考分值）
    question_scores = allocate_question_scores(
        ordered, req.total_score,
        type_distribution=req.type_distribution if not req.question_ids else _infer_distribution(ordered),
        difficulty_distribution=req.difficulty_distribution if not req.question_ids else _infer_diff_distribution(ordered),
    )

    exam = Exam(
        title=req.title,
        description=req.description,
        question_ids=question_ids,
        question_scores=question_scores,
        total_score=req.total_score,
        duration_minutes=req.duration_minutes,
        owner_id=current_user.id,
    )
    db.add(exam)
    db.commit()
    db.refresh(exam)

    return ExamDetailResponse(
        id=exam.id,
        title=exam.title,
        description=exam.description,
        question_ids=exam.question_ids,
        question_scores=exam.question_scores,
        total_score=exam.total_score,
        duration_minutes=exam.duration_minutes,
        created_at=exam.created_at,
        questions=[QuestionResponse.from_question(q, current_user.id) for q in ordered],
    )


@router.get("", response_model=ExamListResponse)
def list_exams(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """试卷列表。"""
    exams = db.query(Exam).filter(_exam_visible(current_user)).order_by(Exam.created_at.desc()).all()
    return ExamListResponse(items=exams, total=len(exams))


@router.get("/{exam_id}", response_model=ExamDetailResponse)
def get_exam(
    exam_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """试卷详情（含完整题目列表）。"""
    exam = db.query(Exam).filter(Exam.id == exam_id, _exam_visible(current_user)).first()
    if not exam:
        raise HTTPException(status_code=404, detail="试卷不存在")

    questions = db.query(Question).filter(Question.id.in_(exam.question_ids)).all()
    q_map = {q.id: q for q in questions}
    ordered = [q_map[qid] for qid in exam.question_ids if qid in q_map]

    return ExamDetailResponse(
        id=exam.id,
        title=exam.title,
        description=exam.description,
        question_ids=exam.question_ids,
        question_scores=exam.question_scores or {},
        total_score=exam.total_score,
        duration_minutes=exam.duration_minutes,
        created_at=exam.created_at,
        questions=[QuestionResponse.from_question(q, current_user.id) for q in ordered],
    )


@router.get("/{exam_id}/analyze")
def analyze_exam_endpoint(
    exam_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """AI 分析试卷难度、知识点覆盖、评分、适合的学生群体。"""
    exam = db.query(Exam).filter(Exam.id == exam_id, _exam_visible(current_user)).first()
    if not exam:
        raise HTTPException(status_code=404, detail="试卷不存在")

    questions = db.query(Question).filter(Question.id.in_(exam.question_ids)).all()
    q_map = {q.id: q for q in questions}
    ordered = [q_map[qid] for qid in exam.question_ids if qid in q_map]

    exam_data = {
        "total_score": exam.total_score,
        "duration_minutes": exam.duration_minutes,
        "total_questions": len(ordered),
        "questions": [
            {
                "type": q.type.value if hasattr(q.type, 'value') else str(q.type),
                "difficulty": q.difficulty.value if hasattr(q.difficulty, 'value') else str(q.difficulty),
                "knowledge_points": q.knowledge_points or [],
                "stem": q.stem or "",
                "answer": q.answer or "",
            }
            for q in ordered
        ],
    }

    result = analyze_exam(exam_data)
    return result


def _latex_to_text(s: str) -> str:
    """将 LaTeX 数学公式转为纯文本 Unicode，用于 Word 导出。"""
    import re
    if not s:
        return s
    r = s

    # 去掉 $$...$$ 和 $...$ 定界符，保留内容
    r = re.sub(r'\$\$(.+?)\$\$', r'\1', r, flags=re.DOTALL)
    r = re.sub(r'\$(.+?)\$', r'\1', r)

    # 常用 LaTeX 命令 → Unicode
    replacements = [
        ('\\triangle', '△'), ('\\angle', '∠'), ('\\pi', 'π'),
        ('\\leq', '≤'), ('\\le', '≤'), ('\\geq', '≥'), ('\\ge', '≥'), ('\\neq', '≠'),
        ('\\times', '×'), ('\\div', '÷'), ('\\cdot', '·'),
        ('\\infty', '∞'), ('\\circ', '°'), ('\\pm', '±'),
        ('\\parallel', '∥'), ('\\perp', '⊥'), ('\\approx', '≈'),
        ('\\cong', '≅'), ('\\sim', '∼'), ('\\equiv', '≡'),
        ('\\Rightarrow', '⇒'), ('\\Leftrightarrow', '⇔'),
        ('\\rightarrow', '→'), ('\\leftarrow', '←'),
        ('\\Longrightarrow', '⟹'), ('\\longrightarrow', '⟶'),
        ('\\square', '□'), ('\\Box', '□'),
        ('\\ldots', '…'), ('\\cdots', '⋯'),
        ('\\emptyset', '∅'), ('\\varnothing', '∅'),
        ('\\in', '∈'), ('\\notin', '∉'), ('\\subset', '⊂'),
        ('\\supset', '⊃'), ('\\cup', '∪'), ('\\cap', '∩'),
        ('\\forall', '∀'), ('\\exists', '∃'),
        ('\\sum', 'Σ'), ('\\prod', 'Π'),
        ('\\alpha', 'α'), ('\\beta', 'β'), ('\\gamma', 'γ'),
        ('\\delta', 'δ'), ('\\theta', 'θ'), ('\\lambda', 'λ'),
        ('\\mu', 'μ'), ('\\sigma', 'σ'), ('\\omega', 'ω'),
        ('\\Omega', 'Ω'), ('\\Delta', 'Δ'),
    ]
    for latex, uni in replacements:
        r = r.replace(latex, uni)

    # \dfrac{a}{b} → \frac{a}{b}，由 _add_styled_paragraph 处理为分数
    r = r.replace('\\dfrac', '\\frac')
    # \frac{a}{b} → 保留，由 _add_styled_paragraph 处理为分数
    # \sqrt{x} 和 \sqrt[n]{x} 保留不转换，由 _add_styled_paragraph 处理为根号EQ域
    # \text{xxx} → xxx
    r = re.sub(r'\\text\{([^}]*)\}', r'\1', r)
    # \textbf{xxx} → xxx
    r = re.sub(r'\\textbf\{([^}]*)\}', r'\1', r)
    # \begin{cases} ... \end{cases} → 去掉环境，保留内容
    r = re.sub(r'\\begin\{cases\}', '', r)
    r = re.sub(r'\\end\{cases\}', '', r)
    # 上标 ^{...} 和 ^x
    r = re.sub(r'\^\{([^}]*)\}', r'^\1', r)
    # 下标 _{...} 和 _x
    r = re.sub(r'_\{([^}]*)\}', r'_\1', r)
    # \left \right 去掉
    r = r.replace('\\left', '').replace('\\right', '')
    r = r.replace('\\,', ' ').replace('\\;', ' ').replace('\\ ', ' ')
    r = r.replace('\\qquad', '  ').replace('\\quad', ' ')
    r = r.replace('\\n', ' ')
    # 转义符号修复
    r = r.replace('\\_', '_')
    r = r.replace('\\%', '%')
    r = r.replace('\\#', '#')
    r = r.replace('\\&', '&')
    # 去掉多余空格
    r = re.sub(r' +', ' ', r).strip()

    return r


def _add_styled_paragraph(doc, text: str, font_size, bold=False, indent=None):
    """添加段落，支持上标、下标和分数格式化。"""
    import re
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    para.paragraph_format.space_before = Pt(2)

    # 第一轮：找出所有特殊块的位置
    BS = chr(92)  # 反斜杠，避免 raw string 转义混乱
    specials = []
    # \frac{num}{den}
    frac_re = re.compile(BS + BS + r'frac\{([^{}]*(?:\{[^}]*\}[^{}]*)*)\}\{([^{}]*(?:\{[^}]*\}[^{}]*)*)\}')
    for m in frac_re.finditer(text):
        specials.append((m.start(), m.end(), 'frac', m.group(1), m.group(2)))
    # \sqrt[n]{x}
    sqrt_n_re = re.compile(BS + BS + r'sqrt\[(\d+)\]\{([^{}]*(?:\{[^}]*\}[^{}]*)*)\}')
    for m in sqrt_n_re.finditer(text):
        specials.append((m.start(), m.end(), 'sqrt_n', m.group(1), m.group(2)))
    # \sqrt{x}
    sqrt_re = re.compile(BS + BS + r'sqrt\{([^{}]*(?:\{[^}]*\}[^{}]*)*)\}')
    for m in sqrt_re.finditer(text):
        specials.append((m.start(), m.end(), 'sqrt', m.group(1), None))
    # ^{...}
    for m in re.finditer(r'\^\{([^}]*)\}', text):
        specials.append((m.start(), m.end(), 'sup_brace', m.group(1), None))
    # _{...}
    for m in re.finditer(r'_\{([^}]*)\}', text):
        specials.append((m.start(), m.end(), 'sub_brace', m.group(1), None))
    # ^x (simple superscript: 数字/字母/Unicode符号)
    for m in re.finditer(r'\^(\d+|[A-Za-z°²³α-ω]+)', text):
        specials.append((m.start(), m.end(), 'sup_simple', m.group(1), None))
    # _x (simple subscript)
    for m in re.finditer(r'_(\d+|[A-Za-z])', text):
        specials.append((m.start(), m.end(), 'sub_simple', m.group(1), None))

    # 排序并去重（按位置排序，同位置取长的）
    specials.sort(key=lambda x: (x[0], -(x[1] - x[0])))
    filtered = []
    last_end = 0
    for s in specials:
        if s[0] >= last_end:
            filtered.append(s)
            last_end = s[1]

    # 第二轮：按顺序输出普通文本 + 特殊块
    pos = 0
    for start, end, kind, arg1, arg2 in filtered:
        # 输出前面的普通文本
        if start > pos:
            run = para.add_run(text[pos:start])
            run.font.size = Pt(font_size)
            run.font.bold = bold
            pos = start

        if kind == 'frac':
            _add_eq_field(para, f' \\f({arg1},{arg2}) ', font_size)
        elif kind == 'sqrt_n':
            _add_eq_field(para, f' \\r({arg1},{arg2}) ', font_size)
        elif kind == 'sqrt':
            _add_eq_field(para, f' \\r(,{arg1}) ', font_size)
        elif kind in ('sup_brace', 'sup_simple'):
            run = para.add_run(arg1)
            run.font.superscript = True
            run.font.size = Pt(font_size * 0.7)
            run.font.bold = bold
        elif kind in ('sub_brace', 'sub_simple'):
            run = para.add_run(arg1)
            run.font.subscript = True
            run.font.size = Pt(font_size * 0.7)
            run.font.bold = bold
        pos = end

    # 输出剩余普通文本
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.size = Pt(font_size)
        run.font.bold = bold

    return para


def _add_eq_field(para, eq_code: str, font_size):
    """在段落中插入 Word EQ 域（分数/根号等）。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = para.add_run()
    run.font.size = Pt(font_size)

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' EQ{eq_code} '
    run._r.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)


@router.get("/{exam_id}/export")
def export_exam(
    exam_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """导出试卷为 Word 文档。"""
    from io import BytesIO

    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    exam = db.query(Exam).filter(Exam.id == exam_id, _exam_visible(current_user)).first()
    if not exam:
        raise HTTPException(status_code=404, detail="试卷不存在")

    questions = db.query(Question).filter(Question.id.in_(exam.question_ids)).all()
    q_map = {q.id: q for q in questions}
    ordered = [q_map[qid] for qid in exam.question_ids if qid in q_map]

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 标题
    title = doc.add_paragraph(exam.title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(18)
    run.font.bold = True

    # 试卷信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"总分：{exam.total_score}分    ")
    info.add_run(f"时长：{exam.duration_minutes}分钟    ")
    info.add_run(f"题量：{len(ordered)}题")

    doc.add_paragraph()  # 空行

    type_labels = {"choice": "选择题", "fill_blank": "填空题", "calculation": "计算题", "proof": "证明题"}

    question_scores = exam.question_scores or {}

    for i, q in enumerate(ordered, 1):
        # 题目（含分值）
        score = question_scores.get(str(i - 1), 0)
        label = type_labels.get(q.type.value if hasattr(q.type, 'value') else str(q.type), q.type)
        score_text = f"（{score}分）" if score else ""
        stem_text = f"{i}. [{label}] {score_text}{_latex_to_text(q.stem or '')}"
        _add_styled_paragraph(doc, stem_text, font_size=12)

        # 选项（选择题）
        if q.options:
            for opt in q.options:
                _add_styled_paragraph(doc, _latex_to_text(opt), font_size=11, indent=1)

    # 答案与解析
    doc.add_page_break()
    answer_title = doc.add_paragraph("参考答案与解析")
    answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answer_title.runs[0].font.size = Pt(16)
    answer_title.runs[0].font.bold = True
    doc.add_paragraph()

    for i, q in enumerate(ordered, 1):
        _add_styled_paragraph(doc, f"{i}. 答案：{_latex_to_text(q.answer or '')}", font_size=11, bold=True)

        if q.answer_analysis:
            _add_styled_paragraph(doc, f"解析：{_latex_to_text(q.answer_analysis)}", font_size=10, indent=1)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    from urllib.parse import quote
    safe_title = exam.title or "试卷"
    filename = f"{safe_title}.docx"
    encoded_filename = quote(filename, safe="")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )


@router.post("/{exam_id}/submit", response_model=ExamResultResponse)
def submit_exam(
    exam_id: int,
    req: ExamSubmitRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """提交答题，自动评分并保存结果。"""
    exam = db.query(Exam).filter(Exam.id == exam_id, _exam_visible(current_user)).first()
    if not exam:
        raise HTTPException(status_code=404, detail="试卷不存在")

    # 服务端超时校验（容忍 2 分钟网络延迟 / 客户端时钟偏差）
    if req.started_at:
        from datetime import datetime as dt, timezone
        try:
            start_dt = dt.fromisoformat(req.started_at)
            elapsed = (dt.now(timezone.utc) - start_dt).total_seconds()
            grace = exam.duration_minutes * 60 + 120
            if elapsed > grace:
                raise HTTPException(status_code=400, detail="考试时间已过，无法提交")
        except ValueError:
            pass  # 时间格式异常则不阻断，信任客户端

    questions = db.query(Question).filter(Question.id.in_(exam.question_ids)).all()
    q_map = {q.id: q for q in questions}
    ordered = [q_map[qid] for qid in exam.question_ids if qid in q_map]

    correct_count = 0
    total_score = 0
    results = {}
    question_scores = exam.question_scores or {}
    # 向后兼容：旧试卷没有每题分值则均分
    if not question_scores and ordered:
        per_q = exam.total_score // len(ordered)
        question_scores = {str(i): per_q for i in range(len(ordered))}

    for i, q in enumerate(ordered):
        user_ans = (req.answers.get(i) or "").strip()
        correct_ans = (q.answer or "").strip()

        is_correct = False
        if q.type.value == "choice":
            is_correct = user_ans.upper()[:1] == correct_ans.upper()[:1]
        elif q.type.value == "proof":
            # 证明题：答案内容必须充分（≥15字），且匹配结论
            user_norm = user_ans.replace(" ", "").lower()
            correct_norm = correct_ans.replace(" ", "").lower()
            if len(correct_norm) <= 5:
                # 短结论（如"成立""全等"）：学生答案需包含结论且 ≥15 字
                is_correct = len(user_ans) >= 15 and correct_norm in user_norm
            else:
                is_correct = len(user_ans) >= 10 and user_norm == correct_norm
        else:
            is_correct = user_ans.replace(" ", "").lower() == correct_ans.replace(" ", "").lower()

        if is_correct:
            correct_count += 1
            total_score += question_scores.get(str(i), 0)
        results[str(i)] = {"user_ans": user_ans, "correct_ans": correct_ans, "correct": is_correct}

    score = total_score

    result = ExamResult(
        exam_id=exam_id,
        user_id=current_user.id,
        score=score,
        total_score=exam.total_score,
        correct_count=correct_count,
        total_count=len(ordered),
        answers=results,
    )
    db.add(result)
    db.commit()
    db.refresh(result)
    return result


@router.get("/{exam_id}/results", response_model=ExamResultListResponse)
def get_exam_results(
    exam_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """查看当前用户在某试卷下的答题记录。"""
    exam = db.query(Exam).filter(Exam.id == exam_id, _exam_visible(current_user)).first()
    if not exam:
        raise HTTPException(status_code=404, detail="试卷不存在")

    results = (
        db.query(ExamResult)
        .filter(ExamResult.exam_id == exam_id, ExamResult.user_id == current_user.id)
        .order_by(ExamResult.created_at.desc())
        .all()
    )
    return ExamResultListResponse(items=results, total=len(results))


@router.get("/results/mine", response_model=ExamResultListResponse)
def get_my_results(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """查看当前用户所有答题记录。"""
    results = (
        db.query(ExamResult)
        .filter(ExamResult.user_id == current_user.id)
        .order_by(ExamResult.created_at.desc())
        .all()
    )
    return ExamResultListResponse(items=results, total=len(results))


@router.delete("/{exam_id}", status_code=204)
def delete_exam(
    exam_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """删除试卷及其关联的答题记录。"""
    exam = db.query(Exam).filter(Exam.id == exam_id, _exam_visible(current_user)).first()
    if not exam:
        raise HTTPException(status_code=404, detail="试卷不存在")
    db.query(ExamResult).filter(ExamResult.exam_id == exam_id).delete()
    db.delete(exam)
    db.commit()
